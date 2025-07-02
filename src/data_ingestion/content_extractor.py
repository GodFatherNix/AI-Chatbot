"""
Content extractor for various file formats found on MOSDAC portal.
"""
import requests
from pathlib import Path
from typing import Dict, Optional, List, Any
import mimetypes
import hashlib

# Document processing imports
try:
    import PyPDF2
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

from bs4 import BeautifulSoup
import json
import re
from dataclasses import dataclass
from urllib.parse import urlparse
import time

from ..utils.logger import get_logger
from ..utils.config import config

logger = get_logger(__name__)

@dataclass
class ExtractedContent:
    """Represents extracted content from a file."""
    source_url: str
    file_path: str
    content_type: str
    title: str
    text_content: str
    metadata: Dict[str, Any]
    tables: List[Dict]
    images: List[Dict]
    links: List[str]
    file_hash: str
    extraction_timestamp: str

class ContentExtractor:
    """
    Extracts content from various file formats.
    Supports PDF, DOCX, XLSX, HTML, and plain text files.
    """
    
    def __init__(self, download_dir: str = "data/raw/downloads"):
        self.download_dir = Path(download_dir)
        self.download_dir.mkdir(parents=True, exist_ok=True)
        
        self.supported_formats = config.get('data_processing.supported_formats', 
                                           ['html', 'pdf', 'docx', 'xlsx', 'txt'])
        
        # Setup session for downloads
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': config.get('portal.user_agent', 'MOSDAC-AI-Bot/1.0')
        })
    
    def _get_file_hash(self, file_path: Path) -> str:
        """Calculate MD5 hash of file."""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _download_file(self, url: str) -> Optional[Path]:
        """Download file from URL."""
        try:
            response = self.session.get(url, timeout=30)
            response.raise_for_status()
            
            # Determine filename
            filename = Path(urlparse(url).path).name
            if not filename or '.' not in filename:
                # Try to get filename from Content-Disposition header
                content_disposition = response.headers.get('content-disposition', '')
                if 'filename=' in content_disposition:
                    filename = content_disposition.split('filename=')[1].strip('"')
                else:
                    # Generate filename based on URL hash
                    filename = f"download_{hashlib.md5(url.encode()).hexdigest()[:8]}"
                    
                    # Try to determine extension from content-type
                    content_type = response.headers.get('content-type', '')
                    if 'pdf' in content_type:
                        filename += '.pdf'
                    elif 'word' in content_type or 'docx' in content_type:
                        filename += '.docx'
                    elif 'excel' in content_type or 'xlsx' in content_type:
                        filename += '.xlsx'
                    elif 'html' in content_type:
                        filename += '.html'
                    else:
                        filename += '.txt'
            
            file_path = self.download_dir / filename
            
            with open(file_path, 'wb') as f:
                f.write(response.content)
            
            logger.info(f"Downloaded file: {url} -> {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Error downloading {url}: {e}")
            return None
    
    def _extract_pdf_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 not available. Install with: pip install PyPDF2")
        
        text_content = ""
        metadata = {}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PdfReader(file)
                
                # Extract metadata
                if pdf_reader.metadata:
                    metadata.update({
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'producer': pdf_reader.metadata.get('/Producer', ''),
                        'creation_date': str(pdf_reader.metadata.get('/CreationDate', '')),
                        'modification_date': str(pdf_reader.metadata.get('/ModDate', ''))
                    })
                
                metadata['page_count'] = len(pdf_reader.pages)
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            text_content = f"Error extracting PDF content: {e}"
        
        return ExtractedContent(
            source_url="",
            file_path=str(file_path),
            content_type='pdf',
            title=metadata.get('title', file_path.stem),
            text_content=text_content.strip(),
            metadata=metadata,
            tables=[],
            images=[],
            links=[],
            file_hash=self._get_file_hash(file_path),
            extraction_timestamp=time.strftime('%Y-%m-%d %H:%M:%S')
        )
    
    def _extract_docx_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not available. Install with: pip install python-docx")
        
        text_content = ""
        tables = []
        metadata = {}
        
        try:
            doc = Document(file_path)
            
            # Extract core properties
            if doc.core_properties:
                metadata.update({
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or '',
                    'keywords': doc.core_properties.keywords or '',
                    'comments': doc.core_properties.comments or '',
                    'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                    'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
                })
            
            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            text_content = '\n'.join(paragraphs)
            
            # Extract tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_id': table_idx,
                        'data': table_data,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0
                    })
            
            metadata['paragraph_count'] = len(paragraphs)
            metadata['table_count'] = len(tables)
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            text_content = f"Error extracting DOCX content: {e}"
        
        return ExtractedContent(
            source_url="",
            file_path=str(file_path),
            content_type='docx',
            title=metadata.get('title', file_path.stem),
            text_content=text_content,
            metadata=metadata,
            tables=tables,
            images=[],
            links=[],
            file_hash=self._get_file_hash(file_path),
            extraction_timestamp=time.strftime('%Y-%m-%d %H:%M:%S')
        )
    
    def _extract_excel_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from Excel file."""
        if not EXCEL_AVAILABLE:
            raise ImportError("openpyxl not available. Install with: pip install openpyxl pandas")
        
        text_content = ""
        tables = []
        metadata = {}
        
        try:
            # Load workbook
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            # Extract workbook properties
            props = workbook.properties
            metadata.update({
                'title': props.title or '',
                'creator': props.creator or '',
                'description': props.description or '',
                'subject': props.subject or '',
                'created': str(props.created) if props.created else '',
                'modified': str(props.modified) if props.modified else '',
                'sheet_count': len(workbook.sheetnames),
                'sheet_names': workbook.sheetnames
            })
            
            # Process each worksheet
            for sheet_idx, sheet_name in enumerate(workbook.sheetnames):
                sheet = workbook[sheet_name]
                
                # Convert sheet to DataFrame for easier processing
                data = []
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):  # Skip empty rows
                        data.append([str(cell) if cell is not None else '' for cell in row])
                
                if data:
                    tables.append({
                        'sheet_name': sheet_name,
                        'sheet_index': sheet_idx,
                        'data': data,
                        'rows': len(data),
                        'columns': len(data[0]) if data else 0
                    })
                    
                    # Add sheet content to text
                    text_content += f"\n--- Sheet: {sheet_name} ---\n"
                    for row in data[:10]:  # Include first 10 rows in text content
                        text_content += '\t'.join(row) + '\n'
                    
                    if len(data) > 10:
                        text_content += f"... ({len(data) - 10} more rows)\n"
            
            metadata['total_tables'] = len(tables)
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path}: {e}")
            text_content = f"Error extracting Excel content: {e}"
        
        return ExtractedContent(
            source_url="",
            file_path=str(file_path),
            content_type='xlsx',
            title=metadata.get('title', file_path.stem),
            text_content=text_content,
            metadata=metadata,
            tables=tables,
            images=[],
            links=[],
            file_hash=self._get_file_hash(file_path),
            extraction_timestamp=time.strftime('%Y-%m-%d %H:%M:%S')
        )
    
    def _extract_html_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from HTML file."""
        text_content = ""
        links = []
        images = []
        tables = []
        metadata = {}
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Extract title
            title_tag = soup.find('title')
            title = title_tag.get_text().strip() if title_tag else file_path.stem
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "header", "footer"]):
                script.decompose()
            
            # Extract main text content
            text_content = soup.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = ' '.join(chunk for chunk in chunks if chunk)
            
            # Extract links
            for link in soup.find_all('a', href=True):
                links.append({
                    'url': link['href'],
                    'text': link.get_text().strip(),
                    'title': link.get('title', '')
                })
            
            # Extract images
            for img in soup.find_all('img'):
                images.append({
                    'src': img.get('src', ''),
                    'alt': img.get('alt', ''),
                    'title': img.get('title', '')
                })
            
            # Extract tables
            for table_idx, table in enumerate(soup.find_all('table')):
                table_data = []
                for row in table.find_all('tr'):
                    row_data = [cell.get_text().strip() for cell in row.find_all(['td', 'th'])]
                    if row_data:
                        table_data.append(row_data)
                
                if table_data:
                    tables.append({
                        'table_id': table_idx,
                        'data': table_data,
                        'rows': len(table_data),
                        'columns': len(table_data[0]) if table_data else 0
                    })
            
            # Extract metadata
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                name = tag.get('name') or tag.get('property')
                content = tag.get('content')
                if name and content:
                    metadata[f"meta_{name}"] = content
            
            metadata.update({
                'link_count': len(links),
                'image_count': len(images),
                'table_count': len(tables)
            })
            
        except Exception as e:
            logger.error(f"Error processing HTML {file_path}: {e}")
            text_content = f"Error extracting HTML content: {e}"
        
        return ExtractedContent(
            source_url="",
            file_path=str(file_path),
            content_type='html',
            title=title,
            text_content=text_content,
            metadata=metadata,
            tables=tables,
            images=images,
            links=links,
            file_hash=self._get_file_hash(file_path),
            extraction_timestamp=time.strftime('%Y-%m-%d %H:%M:%S')
        )
    
    def _extract_text_content(self, file_path: Path) -> ExtractedContent:
        """Extract content from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            metadata = {
                'character_count': len(text_content),
                'line_count': len(text_content.splitlines()),
                'word_count': len(text_content.split())
            }
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {e}")
            text_content = f"Error extracting text content: {e}"
            metadata = {}
        
        return ExtractedContent(
            source_url="",
            file_path=str(file_path),
            content_type='txt',
            title=file_path.stem,
            text_content=text_content,
            metadata=metadata,
            tables=[],
            images=[],
            links=[],
            file_hash=self._get_file_hash(file_path),
            extraction_timestamp=time.strftime('%Y-%m-%d %H:%M:%S')
        )
    
    def extract_from_file(self, file_path: Path) -> Optional[ExtractedContent]:
        """Extract content from a local file."""
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        # Determine file type
        mime_type, _ = mimetypes.guess_type(str(file_path))
        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_pdf_content(file_path)
            elif extension in ['.docx', '.doc']:
                return self._extract_docx_content(file_path)
            elif extension in ['.xlsx', '.xls']:
                return self._extract_excel_content(file_path)
            elif extension in ['.html', '.htm']:
                return self._extract_html_content(file_path)
            elif extension == '.txt' or mime_type and 'text' in mime_type:
                return self._extract_text_content(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            return None
    
    def extract_from_url(self, url: str) -> Optional[ExtractedContent]:
        """Download and extract content from URL."""
        file_path = self._download_file(url)
        if not file_path:
            return None
        
        extracted_content = self.extract_from_file(file_path)
        if extracted_content:
            extracted_content.source_url = url
        
        return extracted_content
    
    def batch_extract(self, file_paths: List[Path]) -> List[ExtractedContent]:
        """Extract content from multiple files."""
        results = []
        
        for file_path in file_paths:
            logger.info(f"Processing file: {file_path}")
            extracted = self.extract_from_file(file_path)
            if extracted:
                results.append(extracted)
        
        logger.info(f"Successfully processed {len(results)} out of {len(file_paths)} files")
        return results
    
    def save_extracted_content(self, content: ExtractedContent, output_dir: str = "data/processed/extracted"):
        """Save extracted content to JSON file."""
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        filename = f"extracted_{content.file_hash}_{int(time.time())}.json"
        file_path = output_path / filename
        
        content_dict = {
            'source_url': content.source_url,
            'file_path': content.file_path,
            'content_type': content.content_type,
            'title': content.title,
            'text_content': content.text_content,
            'metadata': content.metadata,
            'tables': content.tables,
            'images': content.images,
            'links': content.links,
            'file_hash': content.file_hash,
            'extraction_timestamp': content.extraction_timestamp
        }
        
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(content_dict, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Saved extracted content to: {file_path}")